"""
PM 导入导出 Tool - Open WebUI Tool
提供数据导入导出操作

支持格式: json / csv / xlsx
- Excel: 每个子类型一个 sheet（如产品架构: modules / features / parameters 三个 sheet）
- CSV: 每个子类型一个 CSV，多子类型时打包为 zip
- JSON: 单文件包含所有数据

文件流返回机制:
- tool 生成文件后，通过 /api/v1/files/ 上传到 openwebui 文件存储
- 返回 markdown 下载链接，聊天界面自动渲染
- 通过 __event_emitter__ 流式反馈"正在生成…"
"""

import csv
import io
import json
import logging
import zipfile
from typing import Optional

import aiohttp
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from pydantic import BaseModel, Field

log = logging.getLogger(__name__)


# 支持的导出/导入格式
SUPPORTED_FORMATS = ('json', 'csv', 'xlsx', 'markdown', 'docx')

# 各 module_type 对应的子类型拆分规则（用于 Excel 多 sheet 生成）
# 不在此表的 module_type 视为单一子类型（一个 sheet）
MODULE_SUBTYPE_RULES = {
    # 产品架构: 按 entry.data.node_type 拆分 modules / features / parameters
    'architecture': {
        'split_field': 'node_type',
        'subtypes': ['module', 'feature', 'parameter'],
        'sheet_names': {'module': 'modules', 'feature': 'features', 'parameter': 'parameters'},
    },
    'product-architecture': {
        'split_field': 'node_type',
        'subtypes': ['module', 'feature', 'parameter'],
        'sheet_names': {'module': 'modules', 'feature': 'features', 'parameter': 'parameters'},
    },
}

# D4: 字段别名映射 — 当 entry 顶层和 data 中都找不到 key 时，按别名回退查找
# 用于对齐前端表格视图列定义（schedule/roadmap 等模块的视图列）
COLUMN_ALIASES = {
    'assignee': ['assignee', 'owner', 'responsible', '负责人', '责任人'],
    'startDate': ['startDate', 'start_date', 'start', '开始时间', '开始日期', '开始'],
    'endDate': ['endDate', 'end_date', 'end', '结束时间', '结束日期', '结束'],
    'progress': ['progress', '进度'],
    'isMilestone': ['isMilestone', 'is_milestone', 'milestone', '里程碑'],
    'versionId': ['versionId', 'version_id', 'version', '版本ID', '版本'],
    'currentVersionNumber': ['currentVersionNumber', 'current_version_number', '当前版本'],
    'nodeType': ['nodeType', 'node_type', '节点类型'],
    'nodeStatus': ['nodeStatus', 'node_status', '节点状态'],
    'dependencies': ['dependencies', 'deps', '依赖'],
    'calendarEventId': ['calendarEventId', 'calendar_event_id'],
    'updatedAt': ['updatedAt', 'updated_at', '更新时间', '更新'],
    'createdAt': ['createdAt', 'created_at', '创建时间', '创建'],
}

# D4: 按 module_type 默认列定义 — 与前端表格视图列严格对齐
# 当调用方未传 columns 时，按 module_type 取默认列
DEFAULT_COLUMNS_BY_MODULE = {
    'schedule': [
        {'key': 'priority', 'label': '优先级'},
        {'key': 'title', 'label': '任务名称'},
        {'key': 'assignee', 'label': '负责人'},
        {'key': 'startDate', 'label': '开始'},
        {'key': 'endDate', 'label': '结束'},
        {'key': 'progress', 'label': '进度'},
        {'key': 'isMilestone', 'label': '里程碑'},
        {'key': 'currentVersionNumber', 'label': '版本'},
        {'key': 'status', 'label': '状态'},
        {'key': 'updatedAt', 'label': '更新'},
    ],
    'roadmap': [
        {'key': 'priority', 'label': '优先级'},
        {'key': 'title', 'label': '标题'},
        {'key': 'nodeType', 'label': '节点类型'},
        {'key': 'currentVersionNumber', 'label': '版本'},
        {'key': 'nodeStatus', 'label': '节点状态'},
        {'key': 'startDate', 'label': '开始'},
        {'key': 'endDate', 'label': '结束'},
        {'key': 'dependencies', 'label': '依赖'},
        {'key': 'status', 'label': '状态'},
        {'key': 'updatedAt', 'label': '更新'},
    ],
    'requirement': [
        {'key': 'priority', 'label': '优先级'},
        {'key': 'title', 'label': '需求标题'},
        {'key': 'currentVersionNumber', 'label': '版本'},
        {'key': 'status', 'label': '状态'},
        {'key': 'content', 'label': '描述'},
        {'key': 'updatedAt', 'label': '更新'},
    ],
    'risk': [
        {'key': 'priority', 'label': '优先级'},
        {'key': 'title', 'label': '风险'},
        {'key': 'currentVersionNumber', 'label': '版本'},
        {'key': 'status', 'label': '状态'},
        {'key': 'content', 'label': '描述'},
        {'key': 'updatedAt', 'label': '更新'},
    ],
    'meeting': [
        {'key': 'priority', 'label': '优先级'},
        {'key': 'title', 'label': '会议'},
        {'key': 'currentVersionNumber', 'label': '版本'},
        {'key': 'startDate', 'label': '开始'},
        {'key': 'endDate', 'label': '结束'},
        {'key': 'status', 'label': '状态'},
        {'key': 'updatedAt', 'label': '更新'},
    ],
}

# 当 module_type 不在上表中时的通用默认列
DEFAULT_FALLBACK_COLUMNS = [
    {'key': 'id', 'label': 'ID'},
    {'key': 'title', 'label': '标题'},
    {'key': 'module_type', 'label': '类型'},
    {'key': 'status', 'label': '状态'},
    {'key': 'priority', 'label': '优先级'},
    {'key': 'content', 'label': '内容'},
]

# 文件扩展名与 MIME 映射
FORMAT_MIME = {
    'json': ('json', 'application/json'),
    'csv': ('csv', 'text/csv'),
    'xlsx': ('xlsx', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'),
    'markdown': ('md', 'text/markdown'),
    'docx': ('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
}

# PM module_type 在工作台中的友好显示名（用于下载文件命名）
MODULE_DISPLAY_NAMES = {
    'prd': 'PRD',
    'requirement': '需求池',
    'requirement-boundary': '需求边界',
    'roadmap': '路线图',
    'parameter': '参数配置',
    'architecture': '产品架构',
    'product-architecture': '产品架构',
    'prototype': '原型设计',
    'competitor': '竞品分析',
    'spec': 'SPEC规范',
    'flowchart': '流程图',
    'schedule': '项目排期',
    'testcase': '测试用例',
    'risk': '风险分析',
    'meeting': '会议纪要',
    'acceptance': '验收报告',
    'faq': 'FAQ',
}


class Valves(BaseModel):
    """配置参数"""
    pm_api_base_url: str = Field(
        default="http://localhost:8080/api/v1",
        description="PM API 基础 URL"
    )
    pm_api_key: str = Field(
        default="",
        description="PM API 密钥 (可选)"
    )


class Tools:
    def __init__(self):
        self.valves = Valves()

    async def _request(self, method: str, endpoint: str, data: Optional[dict] = None, params: Optional[dict] = None, token: Optional[str] = None) -> dict:
        """发送 HTTP 请求到 PM API（返回 JSON 字典；非 JSON 响应会被包装为 {'raw_text': ...}）"""
        url = f"{self.valves.pm_api_base_url}{endpoint}"
        headers = {"Content-Type": "application/json"}
        if token:
            headers["Authorization"] = f"Bearer {token}"
        elif self.valves.pm_api_key:
            headers["Authorization"] = f"Bearer {self.valves.pm_api_key}"

        try:
            async with aiohttp.ClientSession() as session:
                if method.upper() == "GET":
                    async with session.get(url, headers=headers, params=params) as resp:
                        if resp.status >= 400:
                            return {"error": f"API error {resp.status}", "detail": await resp.text()}
                        try:
                            return await resp.json()
                        except Exception:
                            return {"raw_text": await resp.text()}
                elif method.upper() == "POST":
                    async with session.post(url, headers=headers, json=data) as resp:
                        if resp.status >= 400:
                            return {"error": f"API error {resp.status}", "detail": await resp.text()}
                        try:
                            return await resp.json()
                        except Exception:
                            return {"raw_text": await resp.text()}
                elif method.upper() == "DELETE":
                    async with session.delete(url, headers=headers) as resp:
                        if resp.status >= 400:
                            return {"error": f"API error {resp.status}", "detail": await resp.text()}
                        try:
                            return await resp.json()
                        except Exception:
                            return {"raw_text": await resp.text()}
                else:
                    return {"error": f"Unsupported HTTP method: {method}"}
        except Exception as e:
            log.error(f"PM API request failed: {e}")
            return {"error": str(e)}

    async def _upload_file(self, filename: str, content_bytes: bytes, content_type: str, token: Optional[str] = None) -> dict:
        """
        上传文件到 openwebui 文件存储 (POST /api/v1/files/)
        返回 {id, filename, ...} 或 {error: ...}
        """
        url = f"{self.valves.pm_api_base_url}/files/"
        headers = {}
        if token:
            headers["Authorization"] = f"Bearer {token}"
        elif self.valves.pm_api_key:
            headers["Authorization"] = f"Bearer {self.valves.pm_api_key}"

        try:
            async with aiohttp.ClientSession() as session:
                form = aiohttp.FormData()
                form.add_field('file', content_bytes, filename=filename, content_type=content_type)
                form.add_field('metadata', json.dumps({"name": filename, "content_type": content_type}))
                form.add_field('process', 'false')
                async with session.post(url, headers=headers, data=form) as resp:
                    if resp.status >= 400:
                        return {"error": f"Upload failed {resp.status}", "detail": await resp.text()}
                    return await resp.json()
        except Exception as e:
            log.error(f"File upload failed: {e}")
            return {"error": str(e)}

    def _build_download_link(self, file_id: str, filename: str) -> str:
        """构造 markdown 下载链接，聊天界面会渲染为可点击的下载按钮"""
        return f"📥 [{filename}](/api/v1/files/{file_id}/content?attachment=true)"

    # ===== HTML → Markdown / 纯文本转换 =====
    # 用 bs4 手写转换，避免引入 markdownify 等新依赖。
    # 仅覆盖富文本编辑器常见的 HTML 标签（h1-h6/p/ul/ol/li/table/pre/code/strong/em/a）。
    def _html_to_markdown(self, html: str) -> str:
        """将 HTML 转换为 Markdown 文本。"""
        from bs4 import BeautifulSoup
        if not html:
            return ''
        soup = BeautifulSoup(html, 'html.parser')

        def convert(node) -> str:
            name = (node.name or '').lower()
            if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                level = int(name[1])
                inner = ''.join(convert(c) for c in node.children)
                return f"\n\n{'#' * level} {inner.strip()}\n\n"
            if name == 'p':
                inner = ''.join(convert(c) for c in node.children)
                return f"\n\n{inner.strip()}\n\n"
            if name in ('strong', 'b'):
                inner = ''.join(convert(c) for c in node.children)
                return f"**{inner}**"
            if name in ('em', 'i'):
                inner = ''.join(convert(c) for c in node.children)
                return f"*{inner}*"
            if name == 'code':
                if node.parent and (node.parent.name or '').lower() == 'pre':
                    return node.get_text()
                return f"`{node.get_text()}`"
            if name == 'pre':
                return f"\n\n```\n{node.get_text()}\n```\n\n"
            if name == 'a':
                href = node.get('href', '')
                inner = node.get_text()
                return f"[{inner}]({href})" if href else inner
            if name == 'br':
                return '\n'
            if name == 'hr':
                return '\n\n---\n\n'
            if name in ('ul', 'ol'):
                items = []
                for i, li in enumerate(node.find_all('li', recursive=False)):
                    prefix = f"{i + 1}." if name == 'ol' else '-'
                    items.append(f"{prefix} {li.get_text(strip=True)}")
                return '\n\n' + '\n'.join(items) + '\n\n'
            if name == 'table':
                rows = node.find_all('tr')
                if not rows:
                    return ''
                lines = []
                for idx, tr in enumerate(rows):
                    cells = tr.find_all(['th', 'td'])
                    line = '| ' + ' | '.join(c.get_text(strip=True) for c in cells) + ' |'
                    lines.append(line)
                    if idx == 0:
                        lines.append('| ' + ' | '.join('---' for _ in cells) + ' |')
                return '\n\n' + '\n'.join(lines) + '\n\n'
            # 默认：递归子节点
            if node.children:
                return ''.join(convert(c) for c in node.children)
            return node.get_text() if hasattr(node, 'get_text') else str(node)

        md = convert(soup)
        # 折叠多余空行
        lines = [ln.rstrip() for ln in md.split('\n')]
        result = '\n'.join(lines).strip()
        # 把连续 3+ 个换行压缩为 2 个
        while '\n\n\n' in result:
            result = result.replace('\n\n\n', '\n\n')
        return result

    def _html_to_plain_text(self, html: str) -> str:
        """将 HTML 转换为纯文本（用于 docx / 摘要展示）。"""
        from bs4 import BeautifulSoup
        if not html:
            return ''
        return BeautifulSoup(html, 'html.parser').get_text('\n').strip()

    def _resolve_column_value(self, entry: dict, key: str) -> str:
        """D4: 按别名映射解析字段值。

        查找顺序：
        1. entry 顶层（key 本身）
        2. entry.data（key 本身）
        3. entry 顶层（按 COLUMN_ALIASES[key] 中的别名）
        4. entry.data（按 COLUMN_ALIASES[key] 中的别名）
        5. 兜底返回 ''

        特殊处理：
        - content 字段自动走 _html_to_plain_text 转纯文本
        - data.xxx 形式的 key 直接从 data 中取
        - 数组/对象转 JSON 字符串
        """
        data = entry.get('data') or {}

        # 特殊：data.xxx 形式
        if key.startswith('data.'):
            val = data.get(key[5:], None)
        elif key == 'content':
            val = self._html_to_plain_text(entry.get('content', '') or '')
            return val
        else:
            # 1. 顶层 key 本身
            val = entry.get(key, None)
            # 2. data 中 key 本身
            if val is None or val == '':
                val = data.get(key, None) if isinstance(data, dict) else None
            # 3. 走别名
            if val is None or val == '':
                aliases = COLUMN_ALIASES.get(key, [])
                for alias in aliases:
                    if alias == key:
                        continue
                    top_val = entry.get(alias, None)
                    if top_val not in (None, ''):
                        val = top_val
                        break
                    data_val = data.get(alias, None) if isinstance(data, dict) else None
                    if data_val not in (None, ''):
                        val = data_val
                        break

        # 数组/对象转 JSON 字符串
        if isinstance(val, (list, dict)):
            val = json.dumps(val, ensure_ascii=False)
        # 兜底空值
        if val is None:
            val = ''
        return val

    def _entry_to_row(self, entry: dict, columns: Optional[list] = None) -> dict:
        """将 entry 扁平化为表格行（用于 CSV / Excel）— D4: 完整对齐表格视图列。

        Args:
            entry: PM entry dict
            columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]。
                     若提供则按 columns 顺序输出对应字段；
                     若为 None 则按 module_type 取 DEFAULT_COLUMNS_BY_MODULE 默认列。

        字段 key 解析规则（_resolve_column_value）：
            1. entry 顶层（key 本身）
            2. entry.data（key 本身）
            3. 别名映射（COLUMN_ALIASES）
            4. 兜底 ''
            - content 字段自动走 _html_to_plain_text 转纯文本
            - data.* 嵌套字段直接从 data 中取
        """
        # D4: 未传 columns 时按 module_type 取默认列
        if not columns:
            module_type = entry.get('module_type', '') or ''
            columns = DEFAULT_COLUMNS_BY_MODULE.get(module_type) or DEFAULT_FALLBACK_COLUMNS

        row: dict = {}
        for col in columns:
            key = col.get('key') if isinstance(col, dict) else str(col)
            if not key:
                continue
            row[key] = self._resolve_column_value(entry, key)
        return row

    def _split_entries_by_subtype(self, entries: list, module_type: str, columns: Optional[list] = None) -> dict:
        """
        按 module_type 的子类型规则拆分 entries。
        返回 {subtype_name: [rows]}。若无拆分规则，返回 {'default': [rows]}。
        """
        rows = [self._entry_to_row(e, columns) for e in entries]
        rule = MODULE_SUBTYPE_RULES.get(module_type)
        if not rule:
            return {'default': rows}

        # 按 entry.data[split_field] 拆分
        split_field = rule['split_field']
        sheet_names = rule['sheet_names']
        grouped: dict = {sheet_names[st]: [] for st in rule['subtypes']}
        for entry, row in zip(entries, rows):
            data = entry.get('data') or {}
            subtype = data.get(split_field, '')
            # 兜底：尝试从 content 解析 nodes 列表
            if not subtype and isinstance(entry.get('content'), str):
                try:
                    parsed = json.loads(entry['content'])
                    if isinstance(parsed, dict) and 'nodes' in parsed:
                        for node in parsed.get('nodes', []):
                            nt = node.get('type', '') or node.get('node_type', '')
                            if nt in rule['subtypes']:
                                node_row = dict(row)
                                node_row['title'] = node.get('name', node.get('title', node_row['title']))
                                node_row['data_json'] = json.dumps(node, ensure_ascii=False)
                                grouped.setdefault(sheet_names[nt], []).append(node_row)
                        continue
                except Exception:
                    pass
            if subtype in sheet_names:
                grouped[sheet_names[subtype]].append(row)
            else:
                grouped.setdefault('other', []).append(row)
        return grouped

    def _generate_xlsx_bytes(self, sheets: dict, columns: Optional[list] = None) -> bytes:
        """
        生成 Excel 文件字节流。每个 sheet 一组数据。
        sheets: {sheet_name: [row_dict, ...]}

        样式增强：
        - 表头行：加粗 + 浅蓝背景 + 居中
        - 冻结首行
        - 自动筛选
        - 列宽自适应（content 列固定 40 + wrap_text）
        - 所有单元格 thin border

        Args:
            columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]。
                     若提供则按 columns 顺序与 label 输出表头；
                     若为 None 则使用每行 keys 的并集作为表头（向后兼容，label=key）。
        """
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = Workbook()
        # 移除默认 sheet
        wb.remove(wb.active)

        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        thin_side = Side(border_style='thin', color='D1D5DB')
        cell_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
        cell_align = Alignment(vertical='top', wrap_text=True)

        for sheet_name, rows in sheets.items():
            # sheet 名最长 31 字符
            safe_name = sheet_name[:31] if len(sheet_name) > 31 else sheet_name
            ws = wb.create_sheet(title=safe_name)
            if not rows:
                ws.append(['(empty)'])
                continue
            # D5: 优先用 columns 的 label 作为表头；无 columns 时取 keys 并集（向后兼容）
            if columns:
                col_pairs = self._resolve_column_labels(columns, [])
                # 补齐 rows 中存在但 columns 未声明的 key（如 data_json）
                declared_keys = {k for k, _ in col_pairs}
                extra_keys = []
                for r in rows:
                    for k in r.keys():
                        if k not in declared_keys and k not in extra_keys:
                            extra_keys.append(k)
                col_pairs = col_pairs + [(k, k) for k in extra_keys]
            else:
                all_keys = []
                seen = set()
                for r in rows:
                    for k in r.keys():
                        if k not in seen:
                            seen.add(k)
                            all_keys.append(k)
                col_pairs = [(k, k) for k in all_keys]
            keys = [k for k, _ in col_pairs]
            labels = [lab for _, lab in col_pairs]
            ws.append(labels)
            for r in rows:
                ws.append([r.get(k, '') for k in keys])

            # 表头样式
            for col_idx in range(1, len(labels) + 1):
                cell = ws.cell(row=1, column=col_idx)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = cell_border

            # 数据区样式 + 列宽
            for col_idx, k in enumerate(keys, start=1):
                col_letter = get_column_letter(col_idx)
                # 列宽：content 列固定 40，其他按内容自适应
                if k == 'content' or k == 'data_json':
                    ws.column_dimensions[col_letter].width = 40
                else:
                    cell_lens = [len(str(k))] + [min(len(str(r.get(k, ''))), 200) for r in rows]
                    max_len = max(cell_lens)
                    ws.column_dimensions[col_letter].width = min(max(max_len + 2, 10), 60)
                # 应用对齐 + 边框到所有数据单元格
                for row_idx in range(2, len(rows) + 2):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.alignment = cell_align
                    cell.border = cell_border

            # 冻结首行
            ws.freeze_panes = 'A2'
            # 自动筛选
            ws.auto_filter.ref = ws.dimensions

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def _resolve_column_labels(self, columns: Optional[list], fallback_keys: list) -> list:
        """统一解析 columns 参数，返回 [(key, label), ...]。

        - 若 columns 提供：按 columns 顺序输出，label 取 col['label'] 或 col 自身。
        - 若 columns 为 None：fallback_keys 作为 key+label（向后兼容）。
        """
        if not columns:
            return [(k, k) for k in fallback_keys]
        result = []
        for col in columns:
            if isinstance(col, dict):
                key = col.get('key') or ''
                label = col.get('label') or key
            else:
                key = str(col)
                label = key
            if key:
                result.append((key, label))
        return result

    def _generate_markdown_bytes(self, sheets: dict, module_type: str, columns: Optional[list] = None) -> bytes:
        """生成 Markdown 字节流。

        对于多 sheet（如产品架构的 modules/features/parameters），
        每个 sheet 一个 H2 章节；每行作为表格行输出。
        content 列走 _html_to_markdown 转换（保留格式）。

        Args:
            columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]。
                     若提供则按 columns 顺序与 label 输出表头；
                     若为 None 则使用每行 rows[0].keys() 作为表头（向后兼容）。
        """
        display_name = MODULE_DISPLAY_NAMES.get(module_type, module_type)
        lines = [f"# {display_name}", ""]
        for sheet_name, rows in sheets.items():
            lines.append(f"## {sheet_name}")
            lines.append("")
            if not rows:
                lines.append("_(无数据)_")
                lines.append("")
                continue
            # 列顺序：columns 优先；否则取 rows[0].keys()
            fallback_keys = list(rows[0].keys())
            col_pairs = self._resolve_column_labels(columns, fallback_keys)
            keys = [k for k, _ in col_pairs]
            labels = [lab for _, lab in col_pairs]
            # 表头（使用 label）
            lines.append('| ' + ' | '.join(labels) + ' |')
            lines.append('| ' + ' | '.join('---' for _ in labels) + ' |')
            for r in rows:
                cells = []
                for k in keys:
                    val = str(r.get(k, '') or '')
                    # content 列走 markdown 转换（简化为单行展示）
                    if k == 'content' and val:
                        val = self._html_to_markdown(val).replace('\n', ' ').replace('|', '\\|')
                    else:
                        val = val.replace('\n', ' ').replace('|', '\\|')
                    cells.append(val)
                lines.append('| ' + ' | '.join(cells) + ' |')
            lines.append("")
        return '\n'.join(lines).encode('utf-8')

    def _generate_docx_bytes(self, sheets: dict, module_type: str, columns: Optional[list] = None) -> bytes:
        """生成 docx 字节流。

        每个 sheet 一个 H1，每行一个 H2 + 表格化的元数据 + content 段落。
        content 列走 _html_to_plain_text 转纯文本（docx 不保留 HTML 标签）。

        Args:
            columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]。
                     若提供则按 columns 顺序与 label 输出元数据；
                     若为 None 则使用行内所有非 content/data_json 字段（向后兼容）。
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError as e:
            raise ImportError(
                "导出 Word 文档需要 python-docx 库。请在后端环境执行: pip install python-docx"
            ) from e

        doc = Document()
        display_name = MODULE_DISPLAY_NAMES.get(module_type, module_type)
        doc.add_heading(display_name, level=0)

        for sheet_name, rows in sheets.items():
            doc.add_heading(sheet_name, level=1)
            if not rows:
                doc.add_paragraph('（无数据）')
                continue
            # 列顺序：columns 优先；否则取 rows[0].keys() 过滤掉 content/data_json
            fallback_keys = [k for k in rows[0].keys() if k not in ('content', 'data_json')]
            col_pairs = self._resolve_column_labels(columns, fallback_keys)
            # content/data_json 若在 columns 中显式声明，仍允许出现在 meta_keys
            meta_pairs = col_pairs  # 全部作为元数据键值对
            for r in rows:
                title = str(r.get('title') or r.get('name') or r.get('id') or 'Untitled')
                doc.add_heading(title, level=2)
                # 元数据列表（按 columns 顺序，使用 label 显示）
                for k, label in meta_pairs:
                    if k in ('content', 'data_json'):
                        continue  # content 单独段落输出；data_json 跳过
                    val = r.get(k)
                    if val:
                        doc.add_paragraph(f"{label}: {val}", style='Normal')
                # content 段落
                content = r.get('content')
                if content:
                    doc.add_paragraph(self._html_to_plain_text(content))
                doc.add_paragraph('')  # 空行分隔
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _generate_csv_string(self, rows: list, columns: Optional[list] = None) -> str:
        """生成单个 CSV 字符串

        Args:
            columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]。
                     若提供则按 columns 顺序与 label 输出表头；
                     若为 None 则使用每行 keys 的并集作为表头（向后兼容，label=key）。
        """
        if not rows:
            return ''
        output = io.StringIO()
        writer = csv.writer(output)
        # D5: 优先用 columns 的 label；无 columns 时取 keys 并集（向后兼容）
        if columns:
            col_pairs = self._resolve_column_labels(columns, [])
            declared_keys = {k for k, _ in col_pairs}
            extra_keys = []
            for r in rows:
                for k in r.keys():
                    if k not in declared_keys and k not in extra_keys:
                        extra_keys.append(k)
            col_pairs = col_pairs + [(k, k) for k in extra_keys]
        else:
            all_keys = []
            seen = set()
            for r in rows:
                for k in r.keys():
                    if k not in seen:
                        seen.add(k)
                        all_keys.append(k)
            col_pairs = [(k, k) for k in all_keys]
        keys = [k for k, _ in col_pairs]
        labels = [lab for _, lab in col_pairs]
        writer.writerow(labels)
        for r in rows:
            writer.writerow([r.get(k, '') for k in keys])
        return output.getvalue()

    def _generate_csv_zip_bytes(self, sheets: dict, columns: Optional[list] = None) -> bytes:
        """生成多 CSV 的 zip 包字节流"""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for sheet_name, rows in sheets.items():
                csv_content = self._generate_csv_string(rows, columns=columns)
                zf.writestr(f"{sheet_name}.csv", csv_content)
        return buf.getvalue()

    def _generate_json_bytes(self, sheets: dict, meta: Optional[dict] = None) -> bytes:
        """生成 JSON 字节流（包含所有 sheets + 元信息）"""
        payload = {
            'meta': meta or {},
            'sheets': sheets,
        }
        return json.dumps(payload, ensure_ascii=False, indent=2).encode('utf-8')

    def _parse_xlsx_to_sheets(self, content_bytes: bytes) -> dict:
        """解析 xlsx 字节流为 {sheet_name: [row_dict, ...]}"""
        wb = load_workbook(io.BytesIO(content_bytes), data_only=True)
        sheets = {}
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                sheets[ws.title] = []
                continue
            headers = [str(h) if h is not None else '' for h in rows[0]]
            sheet_rows = []
            for row in rows[1:]:
                row_dict = {}
                for i, val in enumerate(row):
                    if i < len(headers):
                        row_dict[headers[i]] = '' if val is None else str(val)
                sheet_rows.append(row_dict)
            sheets[ws.title] = sheet_rows
        return sheets

    def _parse_csv_to_rows(self, csv_string: str) -> list:
        """解析单个 CSV 字符串为 row 列表"""
        reader = csv.DictReader(io.StringIO(csv_string))
        return [dict(r) for r in reader]

    def _parse_json_to_sheets(self, json_string: str) -> dict:
        """解析 JSON 字符串为 {sheet_name: [row_dict, ...]}"""
        parsed = json.loads(json_string)
        if isinstance(parsed, dict) and 'sheets' in parsed:
            return parsed['sheets']
        if isinstance(parsed, dict):
            # 单 sheet 容错
            return {'default': [parsed]}
        if isinstance(parsed, list):
            return {'default': parsed}
        return {'default': []}

    def _parse_import_data(self, data: str, fmt: str) -> dict:
        """
        将 import 数据字符串解析为 {sheet_name: [row_dict, ...]} 格式。
        支持:
        - json: {sheets: {...}} 或 [...] 或 {...}
        - csv: 单 CSV 字符串 -> {'default': [...]}
        - xlsx: base64 编码的 xlsx 字节流 -> 多 sheet
        """
        import base64
        if fmt == 'json':
            return self._parse_json_to_sheets(data)
        if fmt == 'csv':
            return {'default': self._parse_csv_to_rows(data)}
        if fmt == 'xlsx':
            try:
                # 兼容 base64 或纯二进制字符串
                try:
                    content_bytes = base64.b64decode(data)
                except Exception:
                    content_bytes = data.encode('utf-8') if isinstance(data, str) else data
                return self._parse_xlsx_to_sheets(content_bytes)
            except Exception as e:
                return {'error': f'xlsx 解析失败: {e}'}
        return {'error': f'不支持的格式: {fmt}'}

    async def _emit_preview(self, content: str, __event_emitter__: callable = None):
        if __event_emitter__:
            await __event_emitter__({"type": "message", "data": {"content": content}})

    def _parse_import_preview(self, data: str, export_format: str) -> str:
        try:
            if export_format == "json":
                parsed = json.loads(data)
                if isinstance(parsed, list):
                    count = len(parsed)
                    columns = list(parsed[0].keys()) if parsed else []
                    sample = parsed[:2]
                elif isinstance(parsed, dict):
                    items = parsed.get("items", parsed.get("entries", [parsed]))
                    count = len(items)
                    columns = list(items[0].keys()) if items else []
                    sample = items[:2]
                else:
                    return f"数据预览: {str(data)[:200]}"
                return f"数据预览:\n- 记录数: {count}\n- 列名: {', '.join(columns)}\n- 示例:\n{json.dumps(sample, ensure_ascii=False, indent=2)}"
            elif export_format == "csv":
                lines = data.strip().split("\n")
                total_lines = len(lines) - 1
                header = lines[0] if lines else ""
                sample_lines = lines[1:3]
                return f"数据预览:\n- 行数: {total_lines}\n- 列名: {header}\n- 示例:\n" + "\n".join(sample_lines)
            else:
                return f"数据预览 ({len(data)} 字符):\n{data[:300]}"
        except Exception:
            return f"数据预览 (原始):\n{data[:300]}"

    async def import_entries(self, project_id: str, module_type: str, format: str, data: str, __event_call__: callable = None, __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        导入条目数据到指定项目（预览后确认导入）

        支持格式: json / csv / xlsx
        - xlsx: 解析后只读取第一个 sheet 或名为 'default' 的 sheet
        - 多 sheet 导入请使用 import_module 方法

        :param project_id: 项目 ID
        :param module_type: 模块类型
        :param format: 导入格式 (json/csv/xlsx)
        :param data: 要导入的数据内容（xlsx 时为 base64 编码）
        :param __event_call__: 事件回调 (可选，用于确认流程)
        :param __event_emitter__: 事件发射器 (可选，用于预览数据)
        :return: 导入结果
        """
        if format not in SUPPORTED_FORMATS:
            return json.dumps({"error": f"不支持的格式: {format}，支持: {SUPPORTED_FORMATS}"}, ensure_ascii=False)

        await self._emit_preview(f"⏳ 正在解析 {format} 数据…", __event_emitter__)

        # 解析数据为 sheets 结构
        sheets = self._parse_import_data(data, format)
        if 'error' in sheets:
            return json.dumps(sheets, ensure_ascii=False)

        # 单条目导入: 优先使用 'default' sheet，否则取第一个 sheet
        if 'default' in sheets:
            rows_to_import = sheets['default']
        else:
            rows_to_import = next(iter(sheets.values()), [])

        # 生成预览
        preview_lines = [f"- 解析到 **{len(rows_to_import)}** 条记录"]
        if rows_to_import:
            columns = list(rows_to_import[0].keys())
            preview_lines.append(f"- 列名: {', '.join(columns)}")
            preview_lines.append("- 示例:")
            preview_lines.append(f"```json\n{json.dumps(rows_to_import[0], ensure_ascii=False, indent=2)[:500]}\n```")
        await self._emit_preview(f"**导入数据预览**\n\n" + "\n".join(preview_lines), __event_emitter__)

        if __event_call__:
            confirm = await __event_call__({
                "type": "confirmation",
                "data": {
                    "title": "确认导入数据",
                    "message": f"预览数据已展示。确认导入到项目 {project_id} 的 {module_type} 模块？此操作可能会覆盖现有数据。"
                }
            })
            if not confirm:
                return json.dumps({"status": "cancelled", "message": "用户取消了导入操作"}, ensure_ascii=False)

        # 将 rows 转回 json 字符串交给后端 import 接口（后端只支持 json/csv）
        # 这里统一转成 json 格式提交
        request_data = {
            "module_type": module_type,
            "format": "json",
            "data": json.dumps(rows_to_import, ensure_ascii=False)
        }
        result = await self._request("POST", f"/pm/projects/{project_id}/entries/import", request_data)
        return json.dumps(result, ensure_ascii=False)

    async def export_entry(self, entry_id: str, format: str = "json", __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        导出指定条目数据

        支持格式:
        - json: 返回 JSON 字符串（包含 entry + version）
        - csv: 生成 CSV 文件并上传，返回下载链接
        - xlsx: 生成 Excel 文件（单 sheet）并上传，返回下载链接
        - markdown: 返回 Markdown 文本

        :param entry_id: 条目 ID
        :param format: 导出格式 (json/csv/xlsx/markdown)
        :param __event_emitter__: 事件发射器 (可选，用于预览导出内容)
        :return: 导出的数据内容或下载链接
        """
        if format not in SUPPORTED_FORMATS and format != 'markdown':
            return json.dumps({"error": f"不支持的格式: {format}，支持: {SUPPORTED_FORMATS + ('markdown',)}"}, ensure_ascii=False)

        await self._emit_preview(f"⏳ 正在导出条目 {entry_id} 为 {format}…", __event_emitter__)

        # markdown 直接走后端原接口
        if format == 'markdown':
            params = {"export_format": 'markdown'}
            result = await self._request("GET", f"/pm/entries/{entry_id}/export", params=params)
            if 'raw_text' in result:
                return result['raw_text']
            return json.dumps(result, ensure_ascii=False)

        # 获取 JSON 格式的完整条目数据
        params = {"export_format": 'json'}
        result = await self._request("GET", f"/pm/entries/{entry_id}/export", params=params)
        if 'error' in result:
            return json.dumps(result, ensure_ascii=False)

        entry = result.get('entry', result)
        rows = [self._entry_to_row(entry)]

        # 生成文件并上传
        token = (__user__ or {}).get('token') if __user__ else None
        filename_base = f"entry_{entry_id[:8] if entry_id else 'unknown'}"

        if format == 'json':
            content_bytes = json.dumps(result, ensure_ascii=False, indent=2).encode('utf-8')
            ext, mime = FORMAT_MIME['json']
            filename = f"{filename_base}.{ext}"
            await self._emit_preview(f"⏳ 正在上传文件…", __event_emitter__)
            upload = await self._upload_file(filename, content_bytes, mime, token=token)
            if 'error' in upload:
                return json.dumps(upload, ensure_ascii=False)
            return self._build_download_link(upload.get('id', ''), filename)

        if format == 'csv':
            csv_string = self._generate_csv_string(rows)
            content_bytes = csv_string.encode('utf-8')
            ext, mime = FORMAT_MIME['csv']
            filename = f"{filename_base}.{ext}"
            await self._emit_preview(f"⏳ 正在上传文件…", __event_emitter__)
            upload = await self._upload_file(filename, content_bytes, mime, token=token)
            if 'error' in upload:
                return json.dumps(upload, ensure_ascii=False)
            return self._build_download_link(upload.get('id', ''), filename)

        if format == 'xlsx':
            sheets = {entry.get('module_type', 'entry'): rows}
            content_bytes = self._generate_xlsx_bytes(sheets)
            ext, mime = FORMAT_MIME['xlsx']
            filename = f"{filename_base}.{ext}"
            await self._emit_preview(f"⏳ 正在上传文件…", __event_emitter__)
            upload = await self._upload_file(filename, content_bytes, mime, token=token)
            if 'error' in upload:
                return json.dumps(upload, ensure_ascii=False)
            return self._build_download_link(upload.get('id', ''), filename)

        return json.dumps(result, ensure_ascii=False)

    async def export_module(self, project_id: str, module_type: str, format: str = "xlsx", columns: Optional[list] = None, __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        按模块批量导出（一个 module_type 下所有条目）

        - Excel: 每个子类型一个 sheet（如产品架构: modules / features / parameters）
        - CSV: 多子类型时打包为 zip，单子类型为单 CSV
        - JSON: 单文件包含所有 sheets
        - Markdown / DOCX: 富文本导出，content 列保留格式
        - 关系模块（module_type='relation'）: 走 /pm/projects/{id}/relations 接口，导出双向关系

        :param project_id: 项目 ID
        :param module_type: 模块类型 (如 architecture / requirement / prd / relation 等)
        :param format: 导出格式 (json/csv/xlsx/markdown/docx)
        :param columns: 可选，形如 [{"key": "title", "label": "标题"}, ...]，
                        控制导出列顺序与表头标签；为 None 时使用默认列
        :param __event_emitter__: 事件发射器
        :return: 下载链接（markdown）
        """
        if format not in SUPPORTED_FORMATS:
            return json.dumps({"error": f"不支持的格式: {format}，支持: {SUPPORTED_FORMATS}"}, ensure_ascii=False)

        display_name = MODULE_DISPLAY_NAMES.get(module_type, module_type)
        await self._emit_preview(f"⏳ 正在从 PM 工作台拉取「{display_name}」模块数据…", __event_emitter__)

        token = (__user__ or {}).get('token') if __user__ else None

        # 关系模块特殊处理：走 relations 接口，导出双向关系
        if module_type == 'relation' or module_type == 'relations':
            return await self._export_relations(project_id, format, token, __event_emitter__)

        # 拉取该 module_type 下所有条目
        entries_result = await self._request(
            "GET",
            f"/pm/projects/{project_id}/entries",
            params={"module_type": module_type},
            token=token,
        )

        # 兼容 list / dict 两种返回
        if isinstance(entries_result, dict) and 'error' in entries_result:
            return json.dumps(entries_result, ensure_ascii=False)
        entries = entries_result if isinstance(entries_result, list) else entries_result.get('items', entries_result.get('data', []))

        if not entries:
            return json.dumps({"status": "no_data", "message": f"项目 {project_id} 的「{display_name}」模块没有条目可导出"}, ensure_ascii=False)

        await self._emit_preview(f"✅ 拉取到 {len(entries)} 条记录，正在生成 {format} 文件…", __event_emitter__)

        # 按 module_type 子类型规则拆分（透传 columns 控制列顺序）
        sheets = self._split_entries_by_subtype(entries, module_type, columns)

        # 生成文件
        filename_base = f"{module_type}_{project_id[:8] if project_id else 'unknown'}"
        meta = {
            'project_id': project_id,
            'module_type': module_type,
            'display_name': display_name,
            'total_entries': len(entries),
            'sheets': {k: len(v) for k, v in sheets.items()},
        }

        if format == 'xlsx':
            content_bytes = self._generate_xlsx_bytes(sheets)
            ext, mime = FORMAT_MIME['xlsx']
            filename = f"{display_name}_{filename_base}.{ext}"
        elif format == 'csv':
            # 多 sheet 时打包 zip；单 sheet 直接返回 csv
            if len(sheets) == 1:
                single_rows = next(iter(sheets.values()))
                content_bytes = self._generate_csv_string(single_rows).encode('utf-8')
                ext, mime = FORMAT_MIME['csv']
            else:
                content_bytes = self._generate_csv_zip_bytes(sheets)
                ext, mime = 'zip', 'application/zip'
            filename = f"{display_name}_{filename_base}.{ext}"
        elif format == 'markdown':
            content_bytes = self._generate_markdown_bytes(sheets, module_type, columns=columns)
            ext, mime = FORMAT_MIME['markdown']
            filename = f"{display_name}_{filename_base}.{ext}"
        elif format == 'docx':
            content_bytes = self._generate_docx_bytes(sheets, module_type, columns=columns)
            ext, mime = FORMAT_MIME['docx']
            filename = f"{display_name}_{filename_base}.{ext}"
        else:  # json
            content_bytes = self._generate_json_bytes(sheets, meta=meta)
            ext, mime = FORMAT_MIME['json']
            filename = f"{display_name}_{filename_base}.{ext}"

        await self._emit_preview(f"⏳ 正在上传 {filename}…", __event_emitter__)
        upload = await self._upload_file(filename, content_bytes, mime, token=token)
        if 'error' in upload:
            return json.dumps(upload, ensure_ascii=False)

        link = self._build_download_link(upload.get('id', ''), filename)
        summary = (
            f"✅ 导出完成\n\n"
            f"- 模块: {display_name} (`{module_type}`)\n"
            f"- 总条目数: {len(entries)}\n"
            f"- 格式: {format}\n"
            f"- 文件分 sheet: {', '.join([f'{k}({len(v)})' for k, v in sheets.items()])}\n\n"
            f"{link}"
        )
        return summary

    async def import_module(self, project_id: str, module_type: str, format: str, data: str, __event_call__: callable = None, __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        按模块批量导入（支持多 sheet / 多 CSV zip）

        - xlsx: 解析所有 sheet，按 sheet 名匹配子类型，逐 sheet 批量导入
        - csv: 单 CSV 直接导入；zip 包内多个 CSV 按文件名匹配子类型
        - json: {sheets: {...}} 多 sheet 导入；[...] 单批导入
        - 关系模块（module_type='relation'）: 走专用 _import_relations，自动处理双向关系

        :param project_id: 项目 ID
        :param module_type: 模块类型
        :param format: 导入格式 (json/csv/xlsx)
        :param data: 数据内容（xlsx/zip 时为 base64 编码）
        :param __event_call__: 事件回调 (确认流程)
        :param __event_emitter__: 事件发射器
        :return: 导入结果汇总
        """
        if format not in SUPPORTED_FORMATS:
            return json.dumps({"error": f"不支持的格式: {format}，支持: {SUPPORTED_FORMATS}"}, ensure_ascii=False)

        # 关系模块特殊处理
        if module_type == 'relation' or module_type == 'relations':
            token = (__user__ or {}).get('token') if __user__ else None
            return await self._import_relations(project_id, format, data, token, __event_call__, __event_emitter__)

        display_name = MODULE_DISPLAY_NAMES.get(module_type, module_type)
        await self._emit_preview(f"⏳ 正在解析 {format} 数据…", __event_emitter__)

        # 对于 zip 格式的 CSV，先解包成多 sheet
        if format == 'csv' and data.startswith('PK'):  # zip magic bytes base64
            import base64
            try:
                content_bytes = base64.b64decode(data)
                sheets = self._extract_csv_zip_to_sheets(content_bytes)
            except Exception as e:
                return json.dumps({"error": f"zip 解析失败: {e}"}, ensure_ascii=False)
        else:
            sheets = self._parse_import_data(data, format)

        if 'error' in sheets:
            return json.dumps(sheets, ensure_ascii=False)

        # 预览
        total_rows = sum(len(v) for v in sheets.values())
        preview_lines = [f"- 总记录数: **{total_rows}**", f"- 分 sheet:"]
        for sn, rows in sheets.items():
            preview_lines.append(f"  - `{sn}`: {len(rows)} 条")
        if total_rows > 0:
            first_sheet = next(iter(sheets.values()), [])
            if first_sheet:
                preview_lines.append(f"- 列名: {', '.join(list(first_sheet[0].keys())[:8])}")
        await self._emit_preview(f"**导入数据预览（{display_name}）**\n\n" + "\n".join(preview_lines), __event_emitter__)

        if __event_call__:
            confirm = await __event_call__({
                "type": "confirmation",
                "data": {
                    "title": "确认批量导入",
                    "message": f"将向项目 {project_id} 的「{display_name}」模块导入 {total_rows} 条记录。确认？"
                }
            })
            if not confirm:
                return json.dumps({"status": "cancelled", "message": "用户取消了导入操作"}, ensure_ascii=False)

        # 逐 sheet 批量导入
        token = (__user__ or {}).get('token') if __user__ else None
        all_results = {
            'status': 'completed',
            'module_type': module_type,
            'total_imported': 0,
            'total_errors': 0,
            'sheets': [],
        }

        for sheet_name, rows in sheets.items():
            if not rows:
                continue
            await self._emit_preview(f"⏳ 正在导入 sheet `{sheet_name}` ({len(rows)} 条)…", __event_emitter__)

            # 将行数据转回 entries（识别 data_json 字段）
            entries_to_import = []
            for row in rows:
                entry = {
                    'title': row.get('title', ''),
                    'content': row.get('content', ''),
                    'status': row.get('status', 'draft'),
                    'priority': row.get('priority', ''),
                    'module_type': row.get('module_type') or module_type,
                }
                if row.get('data_json'):
                    try:
                        entry['data'] = json.loads(row['data_json'])
                    except Exception:
                        entry['data'] = {}
                entries_to_import.append(entry)

            request_data = {
                "module_type": module_type,
                "format": "json",
                "data": json.dumps(entries_to_import, ensure_ascii=False)
            }
            sheet_result = await self._request(
                "POST",
                f"/pm/projects/{project_id}/entries/import",
                request_data,
                token=token,
            )

            if isinstance(sheet_result, dict) and 'error' in sheet_result:
                all_results['sheets'].append({
                    'sheet': sheet_name,
                    'imported': 0,
                    'errors': [sheet_result],
                })
                all_results['total_errors'] += 1
            else:
                imported_count = sheet_result.get('imported', 0) if isinstance(sheet_result, dict) else 0
                errors = sheet_result.get('errors', []) if isinstance(sheet_result, dict) else []
                all_results['sheets'].append({
                    'sheet': sheet_name,
                    'imported': imported_count,
                    'errors': errors,
                })
                all_results['total_imported'] += imported_count
                all_results['total_errors'] += len(errors) if isinstance(errors, list) else 0

        summary_lines = [
            f"✅ 批量导入完成（{display_name}）",
            f"",
            f"- 总成功: **{all_results['total_imported']}** 条",
            f"- 总失败: **{all_results['total_errors']}** 条",
            f"- 分 sheet 结果:",
        ]
        for s in all_results['sheets']:
            summary_lines.append(f"  - `{s['sheet']}`: 成功 {s['imported']}, 失败 {len(s.get('errors', []))}")
        await self._emit_preview("\n".join(summary_lines), __event_emitter__)
        return json.dumps(all_results, ensure_ascii=False)

    def _extract_csv_zip_to_sheets(self, zip_bytes: bytes) -> dict:
        """从 zip 字节流中提取多个 CSV，返回 {filename_without_ext: [rows]}"""
        sheets = {}
        with zipfile.ZipFile(io.BytesIO(zip_bytes), 'r') as zf:
            for name in zf.namelist():
                if not name.endswith('.csv'):
                    continue
                csv_content = zf.read(name).decode('utf-8')
                sheet_name = name.rsplit('.', 1)[0]
                sheets[sheet_name] = self._parse_csv_to_rows(csv_content)
        return sheets

    async def _export_relations(self, project_id: str, format: str, token: Optional[str], __event_emitter__: callable = None) -> str:
        """
        关系模块专用导出：调用 /pm/projects/{id}/relations 拉取所有关系，
        导出为 (entity_a_id, entity_b_id, relation_type, confidence, confirmed, direction) 行。
        双向关系（direction='bidirectional'）会同时导出正向和反向两条记录。
        """
        await self._emit_preview(f"⏳ 正在拉取项目 {project_id} 的关系数据…", __event_emitter__)
        result = await self._request("GET", f"/pm/projects/{project_id}/relations", token=token)
        if isinstance(result, dict) and 'error' in result:
            return json.dumps(result, ensure_ascii=False)
        relations = result if isinstance(result, list) else result.get('items', result.get('data', []))

        if not relations:
            return json.dumps({"status": "no_data", "message": f"项目 {project_id} 没有关系可导出"}, ensure_ascii=False)

        await self._emit_preview(f"✅ 拉取到 {len(relations)} 条关系，正在生成 {format} 文件…", __event_emitter__)

        # 双向关系同时导出正向和反向两条记录
        rows = []
        for rel in relations:
            base_row = {
                'id': rel.get('id', ''),
                'project_id': rel.get('project_id', project_id),
                'entity_a_id': rel.get('entity_a_id', ''),
                'entity_b_id': rel.get('entity_b_id', ''),
                'relation_type': rel.get('relation_type', 'references'),
                'confidence': rel.get('confidence', 100),
                'confirmed': rel.get('confirmed', 1),
                'direction': rel.get('direction', 'directional'),
                'version_id': rel.get('version_id', ''),
            }
            rows.append(base_row)
            # 双向关系追加反向行
            if base_row['direction'] == 'bidirectional':
                reverse_row = dict(base_row)
                reverse_row['id'] = f"{base_row['id']}_reverse"
                reverse_row['entity_a_id'] = base_row['entity_b_id']
                reverse_row['entity_b_id'] = base_row['entity_a_id']
                rows.append(reverse_row)

        sheets = {'relations': rows}
        filename_base = f"relations_{project_id[:8] if project_id else 'unknown'}"

        if format == 'xlsx':
            content_bytes = self._generate_xlsx_bytes(sheets)
            ext, mime = FORMAT_MIME['xlsx']
        elif format == 'csv':
            content_bytes = self._generate_csv_string(rows).encode('utf-8')
            ext, mime = FORMAT_MIME['csv']
        else:  # json
            content_bytes = self._generate_json_bytes(sheets, meta={'project_id': project_id, 'module_type': 'relation', 'total': len(rows)})
            ext, mime = FORMAT_MIME['json']

        filename = f"关系_{filename_base}.{ext}"
        await self._emit_preview(f"⏳ 正在上传 {filename}…", __event_emitter__)
        upload = await self._upload_file(filename, content_bytes, mime, token=token)
        if 'error' in upload:
            return json.dumps(upload, ensure_ascii=False)

        link = self._build_download_link(upload.get('id', ''), filename)
        return (
            f"✅ 关系导出完成\n\n"
            f"- 项目: `{project_id}`\n"
            f"- 原始关系数: {len(relations)}\n"
            f"- 导出行数: {len(rows)}（双向关系已展开为正反两条）\n"
            f"- 格式: {format}\n\n"
            f"{link}"
        )

    async def _import_relations(self, project_id: str, format: str, data: str, token: Optional[str], __event_call__: callable = None, __event_emitter__: callable = None) -> str:
        """
        关系模块专用导入：解析数据并逐条调用 /pm/projects/{id}/relations 创建关系。
        双向关系（direction='bidirectional'）会自动创建正向 + 反向两条关系记录。
        """
        await self._emit_preview(f"⏳ 正在解析关系数据…", __event_emitter__)

        if format == 'csv' and data.startswith('PK'):
            import base64
            try:
                content_bytes = base64.b64decode(data)
                sheets = self._extract_csv_zip_to_sheets(content_bytes)
            except Exception as e:
                return json.dumps({"error": f"zip 解析失败: {e}"}, ensure_ascii=False)
        else:
            sheets = self._parse_import_data(data, format)

        if 'error' in sheets:
            return json.dumps(sheets, ensure_ascii=False)

        rows = sheets.get('relations') or sheets.get('default') or []
        if not rows:
            return json.dumps({"status": "no_data", "message": "未解析到关系数据"}, ensure_ascii=False)

        # 预览
        await self._emit_preview(
            f"**关系导入预览**\n\n- 解析到 **{len(rows)}** 条关系\n- 列名: {', '.join(list(rows[0].keys())[:8]) if rows else ''}",
            __event_emitter__
        )

        if __event_call__:
            confirm = await __event_call__({
                "type": "confirmation",
                "data": {
                    "title": "确认导入关系",
                    "message": f"将向项目 {project_id} 导入 {len(rows)} 条关系。确认？"
                }
            })
            if not confirm:
                return json.dumps({"status": "cancelled", "message": "用户取消了导入操作"}, ensure_ascii=False)

        created = 0
        errors = []
        for row in rows:
            direction = row.get('direction', 'directional')
            payload = {
                'entity_a_id': row.get('entity_a_id', ''),
                'entity_b_id': row.get('entity_b_id', ''),
                'relation_type': row.get('relation_type', 'references'),
                'confidence': int(row.get('confidence', 100) or 100),
                'confirmed': int(row.get('confirmed', 1) or 1),
                'version_id': row.get('version_id') or None,
            }
            try:
                res = await self._request("POST", f"/pm/projects/{project_id}/relations", payload, token=token)
                if isinstance(res, dict) and 'error' in res:
                    errors.append({'row': row, 'error': res})
                else:
                    created += 1
                    # 双向关系自动创建反向
                    if direction == 'bidirectional':
                        reverse_payload = {
                            **payload,
                            'entity_a_id': payload['entity_b_id'],
                            'entity_b_id': payload['entity_a_id'],
                        }
                        rev_res = await self._request("POST", f"/pm/projects/{project_id}/relations", reverse_payload, token=token)
                        if isinstance(rev_res, dict) and 'error' in rev_res:
                            errors.append({'row': row, 'error': 'reverse creation failed', 'detail': rev_res})
                        else:
                            created += 1
            except Exception as e:
                errors.append({'row': row, 'error': str(e)})

        await self._emit_preview(
            f"✅ 关系导入完成\n\n- 成功: **{created}** 条\n- 失败: **{len(errors)}** 条",
            __event_emitter__
        )
        return json.dumps({
            'status': 'completed',
            'module_type': 'relation',
            'total_created': created,
            'total_errors': len(errors),
            'errors': errors[:20],  # 限制错误数量避免响应过大
        }, ensure_ascii=False)

    async def extract_parameters(self, entry_id: str, __user__: dict = None) -> str:
        """
        从条目中提取参数

        :param entry_id: 条目 ID
        :return: 提取的参数结果
        """
        result = await self._request("POST", f"/pm/entries/{entry_id}/extract-parameters")
        return json.dumps(result, ensure_ascii=False)

    async def generate_content(self, entry_id: str, module_type: str, instructions: str = "", __event_call__: callable = None, __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        使用 AI 生成内容（预览后确认写入）

        :param entry_id: 条目 ID
        :param module_type: 模块类型
        :param instructions: 生成指令 (可选)
        :param __event_call__: 事件回调 (可选，用于确认流程)
        :param __event_emitter__: 事件发射器 (可选，用于预览生成内容)
        :return: 生成的内容结果
        """
        if __event_call__:
            confirm = await __event_call__({
                "type": "confirmation",
                "data": {
                    "title": "确认生成内容",
                    "message": f"确定要为条目 {entry_id} 生成 {module_type} 内容吗？"
                }
            })
            if not confirm:
                return json.dumps({"status": "cancelled", "message": "用户取消了生成操作"}, ensure_ascii=False)
        
        req_data = {
            "module_type": module_type,
            "instructions": instructions
        }
        result = await self._request("POST", f"/pm/entries/{entry_id}/generate", req_data)

        if "error" not in result:
            preview_content = result.get("content", json.dumps(result, ensure_ascii=False))
            await self._emit_preview(f"**生成内容预览 ({module_type})**\n\n{preview_content}", __event_emitter__)
            if __event_call__:
                write_confirm = await __event_call__({
                    "type": "confirmation",
                    "data": {
                        "title": "确认写入内容",
                        "message": "以上生成内容将写入项目。确认写入？"
                    }
                })
                if not write_confirm:
                    return json.dumps({"status": "cancelled", "message": "用户取消了写入操作", "preview": result}, ensure_ascii=False)

        return json.dumps(result, ensure_ascii=False)

    def _entry_to_markdown(self, entry: dict) -> str:
        title = entry.get("title", "Untitled")
        module_type = entry.get("module_type", "unknown")
        status = entry.get("status", "draft")
        priority = entry.get("priority", "")
        content = entry.get("content", "")
        entry_data = entry.get("data", {})

        lines = [
            f"# {title}",
            "",
            f"- **类型**: {module_type}",
            f"- **状态**: {status}",
        ]
        if priority:
            lines.append(f"- **优先级**: {priority}")
        lines.append("")

        if content:
            lines.append(content)
            lines.append("")

        if isinstance(entry_data, dict) and entry_data:
            lines.append("## 详细信息")
            lines.append("")
            for key, value in entry_data.items():
                if key in ("node_type", "start_date", "end_date", "dependencies"):
                    lines.append(f"- **{key}**: {value}")
            lines.append("")

        return "\n".join(lines)

    async def sync_entries_to_knowledge(self, project_id: str, knowledge_id: str = None, __event_call__: callable = None, __event_emitter__: callable = None, __user__: dict = None) -> str:
        """
        将项目条目同步到 Knowledge Base（预览后确认）

        :param project_id: 项目 ID
        :param knowledge_id: 目标知识库 ID (可选，不提供则自动创建)
        :param __event_call__: 事件回调 (可选，用于确认流程)
        :param __event_emitter__: 事件发射器 (可选，用于预览)
        :return: 同步结果
        """
        entries_result = await self._request("GET", f"/pm/projects/{project_id}/entries")
        if isinstance(entries_result, dict) and "error" in entries_result:
            return json.dumps(entries_result, ensure_ascii=False)

        entries = entries_result if isinstance(entries_result, list) else entries_result.get("items", entries_result.get("data", []))
        if not entries:
            return json.dumps({"status": "no_data", "message": f"项目 {project_id} 没有条目可同步"}, ensure_ascii=False)

        markdown_files = []
        for entry in entries:
            if isinstance(entry, dict) and entry.get("title"):
                md_content = self._entry_to_markdown(entry)
                markdown_files.append({
                    "filename": f"{entry.get('module_type', 'entry')}_{entry.get('id', 'unknown')}.md",
                    "content": md_content,
                    "entry_id": entry.get("id"),
                    "title": entry.get("title"),
                })

        preview_lines = [f"将同步 **{len(markdown_files)}** 个条目到 Knowledge Base:"]
        for mf in markdown_files[:10]:
            preview_lines.append(f"  - {mf['filename']} ({mf['title']})")
        if len(markdown_files) > 10:
            preview_lines.append(f"  - ... 还有 {len(markdown_files) - 10} 个")
        await self._emit_preview("\n".join(preview_lines), __event_emitter__)

        if __event_call__:
            confirm = await __event_call__({
                "type": "confirmation",
                "data": {
                    "title": "确认同步到 Knowledge Base",
                    "message": f"将同步 {len(markdown_files)} 个条目到知识库。确认？"
                }
            })
            if not confirm:
                return json.dumps({"status": "cancelled", "message": "用户取消了同步操作"}, ensure_ascii=False)

        if not knowledge_id:
            kb_result = await self._request("POST", "/knowledge/create", {
                "name": f"PM Project {project_id}",
                "description": f"PM 项目 {project_id} 的条目知识库，包含需求、PRD、参数等文档"
            })
            if isinstance(kb_result, dict) and "error" in kb_result:
                return json.dumps(kb_result, ensure_ascii=False)
            knowledge_id = kb_result.get("id")
            if not knowledge_id:
                return json.dumps({"error": "Failed to create knowledge base", "detail": kb_result}, ensure_ascii=False)

        synced = []
        errors = []
        for mf in markdown_files:
            try:
                upload_result = await self._request("POST", "/files/", {
                    "filename": mf["filename"],
                    "content": mf["content"]
                })
                file_id = upload_result.get("id")
                if not file_id:
                    errors.append({"file": mf["filename"], "error": "upload failed", "detail": upload_result})
                    continue

                add_result = await self._request("POST", f"/knowledge/{knowledge_id}/file/add", {
                    "file_id": file_id
                })
                if isinstance(add_result, dict) and "error" in add_result:
                    errors.append({"file": mf["filename"], "error": "add to knowledge failed", "detail": add_result})
                else:
                    synced.append(mf["filename"])
            except Exception as e:
                errors.append({"file": mf["filename"], "error": str(e)})

        return json.dumps({
            "status": "completed",
            "knowledge_id": knowledge_id,
            "synced_count": len(synced),
            "error_count": len(errors),
            "synced_files": synced,
            "errors": errors,
        }, ensure_ascii=False)
